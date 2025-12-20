"""
文档解析工具模块
支持解析 PDF、DOCX、DOC 等格式的文档
"""
import os
import hashlib
from pathlib import Path
from typing import Optional, List, Tuple
from urllib.parse import urlparse
import requests
from app.core.logger import get_logger

logger = get_logger("document_parser")


class DocumentParser:
    """文档解析器，支持多种文档格式"""

    # 支持的文档格式
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}

    def __init__(self):
        """初始化文档解析器"""
        pass

    @staticmethod
    def is_url(path: str) -> bool:
        """判断是否为 URL 路径"""
        try:
            result = urlparse(path)
            return result.scheme in ('http', 'https')
        except:
            return False

    @staticmethod
    def is_local_path(path: str) -> bool:
        """判断是否为本地路径"""
        return os.path.exists(path)

    @staticmethod
    def download_file(url: str, save_dir: str = None) -> Optional[str]:
        """
        从 URL 下载文件到本地临时目录

        Args:
            url: 文件 URL
            save_dir: 保存目录，默认为系统临时目录

        Returns:
            本地文件路径，失败返回 None
        """
        try:
            if save_dir is None:
                save_dir = os.path.join(os.path.dirname(__file__), '../../temp_downloads')

            os.makedirs(save_dir, exist_ok=True)

            # 从 URL 提取文件名
            filename = os.path.basename(urlparse(url).path)
            if not filename:
                # 如果无法从 URL 获取文件名，使用 hash 生成
                filename = hashlib.md5(url.encode()).hexdigest()

            local_path = os.path.join(save_dir, filename)

            # 下载文件
            logger.info(f"正在下载文件: {url}")
            response = requests.get(url, timeout=30, stream=True)
            response.raise_for_status()

            with open(local_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)

            logger.info(f"文件下载成功: {local_path}")
            return local_path

        except Exception as e:
            logger.error(f"下载文件失败 {url}: {e}")
            return None

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """从 PDF 文件中提取文本"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            full_text = '\n'.join(text_parts)
            logger.debug(f"PDF 解析成功: {file_path}, 提取 {len(full_text)} 字符")
            return full_text

        except Exception as e:
            logger.error(f"解析 PDF 失败 {file_path}: {e}")
            raise

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """从 DOCX 文件中提取文本"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            full_text = '\n'.join(text_parts)
            logger.debug(f"DOCX 解析成功: {file_path}, 提取 {len(full_text)} 字符")
            return full_text

        except Exception as e:
            logger.error(f"解析 DOCX 失败 {file_path}: {e}")
            raise

    @staticmethod
    def extract_text_from_doc(file_path: str) -> str:
        """
        从 DOC 文件中提取文本
        注意：.doc 是旧版 Word 格式，解析较为复杂
        推荐转换为 .docx 格式
        """
        try:
            # 尝试使用 textract (需要安装 textract 库)
            import textract

            text = textract.process(file_path).decode('utf-8')
            logger.debug(f"DOC 解析成功: {file_path}, 提取 {len(text)} 字符")
            return text

        except ImportError:
            logger.warning(f"textract 未安装，无法解析 .doc 格式。建议转换为 .docx 格式")
            raise ValueError("不支持 .doc 格式，请安装 textract 或转换为 .docx 格式")
        except Exception as e:
            logger.error(f"解析 DOC 失败 {file_path}: {e}")
            raise

    def parse_document(self, file_path: str) -> Tuple[str, dict]:
        """
        解析文档，返回文本内容和元数据

        Args:
            file_path: 文件路径（本地或 URL）

        Returns:
            (文本内容, 元数据字典)
        """
        # 如果是 URL，先下载
        is_url = self.is_url(file_path)
        original_path = file_path

        if is_url:
            file_path = self.download_file(file_path)
            if not file_path:
                raise ValueError(f"无法下载文件: {original_path}")

        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 获取文件扩展名
        ext = Path(file_path).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        # 解析文档
        try:
            if ext == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif ext == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif ext == '.doc':
                text = self.extract_text_from_doc(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {ext}")

            # 构建元数据
            metadata = {
                'filename': os.path.basename(original_path),
                'file_type': ext[1:],  # 去掉点号
                'file_size': os.path.getsize(file_path),
                'source': 'url' if is_url else 'local',
                'original_path': original_path
            }

            # 如果是下载的临时文件，清理
            if is_url:
                try:
                    os.remove(file_path)
                    logger.debug(f"已删除临时文件: {file_path}")
                except:
                    pass

            return text, metadata

        except Exception as e:
            # 清理临时文件
            if is_url and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
            raise

    def scan_directory(self, directory: str, recursive: bool = True) -> List[str]:
        """
        扫描目录，返回所有支持的文档文件路径

        Args:
            directory: 目录路径（本地或 URL）
            recursive: 是否递归扫描子目录

        Returns:
            文件路径列表
        """
        file_list = []

        # 如果是 URL 目录，不支持扫描（需要服务器支持目录列表）
        if self.is_url(directory):
            logger.warning("不支持扫描 URL 目录，请提供具体的文件 URL 列表")
            return file_list

        # 检查目录是否存在
        if not os.path.isdir(directory):
            raise NotADirectoryError(f"不是有效的目录: {directory}")

        # 扫描目录
        if recursive:
            for root, _, files in os.walk(directory):
                for filename in files:
                    ext = Path(filename).suffix.lower()
                    if ext in self.SUPPORTED_EXTENSIONS:
                        file_list.append(os.path.join(root, filename))
        else:
            for filename in os.listdir(directory):
                file_path = os.path.join(directory, filename)
                if os.path.isfile(file_path):
                    ext = Path(filename).suffix.lower()
                    if ext in self.SUPPORTED_EXTENSIONS:
                        file_list.append(file_path)

        logger.info(f"扫描目录 {directory}: 找到 {len(file_list)} 个文档文件")
        return file_list

    def batch_parse_documents(self, file_paths: List[str]) -> List[Tuple[str, str, dict]]:
        """
        批量解析文档

        Args:
            file_paths: 文件路径列表

        Returns:
            [(文件路径, 文本内容, 元数据), ...]
        """
        results = []

        for i, file_path in enumerate(file_paths):
            try:
                logger.info(f"正在解析文档 {i+1}/{len(file_paths)}: {file_path}")
                text, metadata = self.parse_document(file_path)
                results.append((file_path, text, metadata))
            except Exception as e:
                logger.error(f"解析文档失败 {file_path}: {e}")
                # 继续处理其他文档
                continue

        logger.info(f"批量解析完成: 成功 {len(results)}/{len(file_paths)}")
        return results


# 全局单例
document_parser = DocumentParser()
